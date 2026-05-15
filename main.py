import base64
import json
import os
import re
import shutil
from docx import Document


#
def base64_txt_to_json(input_file, output_file=None):
    with open(input_file, "r") as file:
        base64_str = file.read().strip()

    decoded_str = base64.b64decode(base64_str).decode("utf-8")
    json_data = json.loads(decoded_str)

    if output_file:
        with open(output_file, "w") as outfile:
            json.dump(json_data, outfile, indent=4)

    return json_data



def extract_cndag_jobs(json_data):
    jobs = []

    for section in json_data:
        for content in section.get("content", []):
            for field in content.get("fields", []):
                for condition in field.get("conditions", []):
                    func = condition.get("function", {})
                    for arg in func.get("argList", []):
                        ds = arg.get("dataSource", {})

                        if ds.get("type") == "CNDAG PIPELINE":
                            jobs.append({
                                "section_name": section["sectionName"],
                                "section_search": ds.get(
                                    "sectionSearch",
                                    section["sectionName"]
                                ),
                                "ds_id": ds.get("dsId")
                            })

    return jobs



def copy_reference_doc(reference_doc_path, output_doc_path):
    shutil.copy(reference_doc_path, output_doc_path)



def keep_only_section(doc_path: str, section_search: str):
    doc = Document(doc_path)
    body = doc.element.body
    blocks = list(body)

    para_map = {p._element: p for p in doc.paragraphs}

    start_idx = None
    start_level = None

    # Find section start
    for i, block in enumerate(blocks):
        para = para_map.get(block)
        if para:
            text = para.text.strip()
            if section_search.lower() in text.lower():
                if para.style.name.startswith("Heading"):
                    start_idx = i
                    start_level = int(para.style.name.split()[-1])
                    break

    if start_idx is None:
        doc.save(doc_path)
        return

    # Find section end
    end_idx = len(blocks)

    for i in range(start_idx + 1, len(blocks)):
        para = para_map.get(blocks[i])
        if para and para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1])
            if level <= start_level:
                end_idx = i
                break

    # Remove everything
    for block in blocks:
        body.remove(block)

    # Keep only section
    for block in blocks[start_idx:end_idx]:
        body.append(block)

    # Remove numbering
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            para.text = re.sub(r"^\d+(\.\d+)*\s*", "", para.text.strip())

    doc.save(doc_path)



def generate_all_cndag_subdocs(json_data, reference_doc_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)

    jobs = extract_cndag_jobs(json_data)
    results = []

    for job in jobs:
        section = job["section_search"]
        safe_name = section.replace(" ", "_")

        output_path = os.path.join(
            output_dir,
            f"CNDAG_{safe_name}.docx"
        )

        # Copy reference
        copy_reference_doc(reference_doc_path, output_path)

        # Keep only that section
        keep_only_section(output_path, section)

        results.append({
            "section": section,
            "file_path": output_path,
            "ds_id": job["ds_id"]
        })

    return results



def create_section_mapping(original_json, generated_docs):
    mappings = []

    for section in original_json:
        section_name = section.get("sectionName")
        section_number = section.get("sectionNumber")

        for content in section.get("content", []):
            for field in content.get("fields", []):
                jinja_tag = field.get("jinjaTag")

                content_type = "GENAI"
                section_search = section_name

                for condition in field.get("conditions", []):
                    func = condition.get("function", {})
                    for arg in func.get("argList", []):
                        ds = arg.get("dataSource", {})

                        if ds.get("type") == "CNDAG PIPELINE":
                            content_type = "CNDAG PIPELINE"
                            section_search = ds.get("sectionSearch", section_name)

                subdoc = next(
                    (d for d in generated_docs if d["section"] == section_search),
                    None
                )

                mappings.append({
                    "sectionName": section_name,
                    "sectionNumber": section_number,
                    "jinjaTag": jinja_tag,
                    "contentType": content_type,
                    "subDocxPath": subdoc["file_path"] if content_type == "CNDAG PIPELINE" else None,
                    "generatedContent": "" if content_type == "CNDAG PIPELINE" else "random text"
                })

    return mappings



def generate_ordered_final_doc(mapping_json, reference_doc_path, output_doc_path):
    # Start from original doc
    copy_reference_doc(reference_doc_path, output_doc_path)

    doc = Document(output_doc_path)
    body = doc.element.body
    blocks = list(body)

    para_map = {p._element: p for p in doc.paragraphs}

    sections_to_keep = []

    for item in mapping_json:
        target = item["sectionName"].lower()

        for i, block in enumerate(blocks):
            para = para_map.get(block)
            if para and para.style.name.startswith("Heading"):
                text = para.text.strip().lower()

                if target in text:
                    start = i
                    level = int(para.style.name.split()[-1])

                    end = len(blocks)
                    for j in range(i + 1, len(blocks)):
                        p2 = para_map.get(blocks[j])
                        if p2 and p2.style.name.startswith("Heading"):
                            lvl = int(p2.style.name.split()[-1])
                            if lvl <= level:
                                end = j
                                break

                    sections_to_keep.append((item, start, end))
                    break

    new_blocks = []

    for item, start, end in sections_to_keep:

        if item["contentType"] == "CNDAG PIPELINE":
            new_blocks.extend(blocks[start:end])

        else:
            new_blocks.append(blocks[start])  # heading only
            para = doc.add_paragraph(item["generatedContent"])
            new_blocks.append(para._element)

    # Replace doc content
    for block in blocks:
        body.remove(block)

    for block in new_blocks:
        body.append(block)

    doc.save(output_doc_path)

    print(f"\n✅ Final ordered doc created: {output_doc_path}")



if __name__ == "__main__":

    json_data = base64_txt_to_json("json_b64.txt")

    # ✅ Subdocs
    generated_subdocs = generate_all_cndag_subdocs(
        json_data=json_data,
        reference_doc_path="PIPELINE_EXECUTION_UPGRADE 11.docx",
        output_dir="output_docs"
    )

    print("\n✅ Individual CNDAG sub-documents created:")
    for d in generated_subdocs:
        print(f"- {d['section']} -> {d['file_path']}")

    # ✅ Mapping
    mapping_json = create_section_mapping(
        original_json=json_data,
        generated_docs=generated_subdocs
    )

    with open("section_mapping.json", "w", encoding="utf-8") as f:
        json.dump(mapping_json, f, indent=4)

    print("\n✅ Section mapping JSON created")

    # ✅ Final doc
    generate_ordered_final_doc(
        mapping_json,
        "PIPELINE_EXECUTION_UPGRADE 11.docx",
        "final_cndag_doc.docx"
    )
